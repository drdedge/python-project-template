#!/usr/bin/env python3
"""
Document Formatter Utility
Converts markdown or HTML content to formatted Word documents with customizable styling.
"""


import os
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import mistune
from htmldocx import HtmlToDocx

class WordDocFormatter:
    """
    Utility class for converting markdown or HTML to formatted Word documents.
    Now ensures that *all* body text—including numbered/bulleted lists—inherits
    the user‑specified font and size by rewriting the Normal style and every
    built‑in list style before any content is added.
    """

    def __init__(
        self,
        font_name: str = "Calibri",
        font_size: int = 10,
        heading_font: str = "Calibri",
        heading_sizes: dict | None = None,
        heading_color: str = "#025B95",
        table_header_bg_color: str = "#025B95",
        custom_table_style: bool = True,
        page_orientation: str = "portrait",
        margins: dict | None = None,
    ) -> None:
        if heading_sizes is None:
            heading_sizes = {"h1": 16, "h2": 14, "h3": 12, "h4": 11}
        if margins is None:
            margins = {"top": 1, "bottom": 1, "left": 1, "right": 1}

        self.font_name = font_name
        self.font_size = font_size
        self.heading_font = heading_font
        self.heading_sizes = heading_sizes
        self.heading_color = heading_color
        self.table_header_bg_color = table_header_bg_color
        self.custom_table_style = custom_table_style
        self.page_orientation = page_orientation.lower()
        self.margins = margins

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _apply_base_paragraph_styles(self, doc: Document) -> None:
        """Set Normal + every built‑in *list* paragraph style to the body font."""
        target_styles: list[str] = [
            "Normal",
            "List Paragraph",
            "List Bullet",
            "List Number",
            "List Continue",
        ]
        for style in doc.styles:
            if (
                style.type == WD_STYLE_TYPE.PARAGRAPH
                and style.name in target_styles or style.name.startswith("List")
            ):
                style.font.name = self.font_name
                style.font.size = Pt(self.font_size)

    def _format_tables(self, doc: Document) -> None:
        """Apply the custom table formatting requested by the user."""
        borders_xml = f"""
        <w:tblBorders {nsdecls('w')}>
            <w:top w:val="single" w:sz="4" w:color="auto"/>
            <w:left w:val="nil"/>
            <w:bottom w:val="single" w:sz="4" w:color="auto"/>
            <w:right w:val="nil"/>
            <w:insideH w:val="single" w:sz="4" w:color="auto"/>
            <w:insideV w:val="nil"/>
        </w:tblBorders>
        """
        margin_xml = f"""
        <w:tcMar {nsdecls('w')}>
            <w:top w:w="70" w:type="dxa"/>
            <w:bottom w:w="70" w:type="dxa"/>
            <w:left w:w="100" w:type="dxa"/>
            <w:right w:w="100" w:type="dxa"/>
        </w:tcMar>
        """
        header_bg = self.table_header_bg_color.lstrip("#")

        for table in doc.tables:
            table._tblPr.append(parse_xml(borders_xml))
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    cell._element.tcPr.append(parse_xml(margin_xml))
                    for paragraph in cell.paragraphs:
                        pf = paragraph.paragraph_format
                        pf.space_after = Pt(0)
                        pf.space_before = Pt(0)
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        for run in paragraph.runs:
                            run.font.name = self.font_name
                            run.font.size = Pt(self.font_size)
                            if row_idx == 0:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                    if row_idx == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="{header_bg}"/>'
                        )
                        cell._tc.get_or_add_tcPr().append(shading)

    def _apply_heading_styles(self, doc: Document) -> None:
        """Apply heading overrides (run-level) to the document."""
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                try:
                    level = int(p.style.name[-1])
                except ValueError:
                    level = 1
                if level <= 4:
                    for run in p.runs:
                        run.font.name = self.heading_font
                        run.font.size = Pt(
                            self.heading_sizes.get(f"h{level}", self.font_size)
                        )
                        run.font.bold = True
                        if self.heading_color:
                            hex_color = self.heading_color.lstrip("#")
                            r, g, b = (
                                int(hex_color[:2], 16),
                                int(hex_color[2:4], 16),
                                int(hex_color[4:6], 16),
                            )
                            run.font.color.rgb = RGBColor(r, g, b)

    def _prepare_html(self, html_content: str) -> str:
        """Ensure HTML content is properly wrapped."""
        html = html_content.strip()
        if html and not html.startswith("<"):
            html = f"<div>{html}</div>"
        if html and not (
            html.startswith("<!DOCTYPE") or html.startswith("<html")
        ):
            html = f"<html><body>{html}</body></html>"
        return html

    def _convert_html_to_docx(
        self,
        html_content: str,
        output_filepath: str | None = None,
        reference_docx_path: str | None = None,
    ):
        """Internal method to convert HTML to DOCX with all formatting applied."""
        try:
            # 1️⃣  Create or load the document
            doc = (
                Document(reference_docx_path)
                if reference_docx_path and os.path.exists(reference_docx_path)
                else Document()
            )

            # 2️⃣  Make sure the default *and* list styles are correct *before* adding content
            self._apply_base_paragraph_styles(doc)

            # 3️⃣  Page layout                                                            
            section = doc.sections[0]
            if self.page_orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )
            section.top_margin = Inches(self.margins["top"])
            section.bottom_margin = Inches(self.margins["bottom"])
            section.left_margin = Inches(self.margins["left"])
            section.right_margin = Inches(self.margins["right"])

            # 4️⃣  HTML → DOCX
            html = self._prepare_html(html_content)
            HtmlToDocx().add_html_to_document(html, doc)

            # 5️⃣  Additional table styling
            if self.custom_table_style:
                self._format_tables(doc)

            # 6️⃣  Heading overrides (run‑level)
            self._apply_heading_styles(doc)

            # 7️⃣  Return or save
            if output_filepath:
                doc.save(output_filepath)
                return True
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        except Exception as exc:
            raise ValueError(f"Document conversion failed: {exc}")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def markdown_to_docx(
        self,
        markdown_text: str,
        output_filepath: str | None = None,
        reference_docx_path: str | None = None,
    ):
        """Convert *markdown_text* to a formatted DOCX with the configured styles."""
        # Convert markdown to HTML first
        html = mistune.create_markdown(plugins=["table"], escape=False)(markdown_text)
        # Use the HTML conversion method
        return self._convert_html_to_docx(html, output_filepath, reference_docx_path)

    def html_to_docx(
        self,
        html_text: str,
        output_filepath: str | None = None,
        reference_docx_path: str | None = None,
    ):
        """Convert *html_text* to a formatted DOCX with the configured styles."""
        return self._convert_html_to_docx(html_text, output_filepath, reference_docx_path)

    def to_docx(
        self,
        content: str,
        output_filepath: str | None = None,
        reference_docx_path: str | None = None,
        input_format: str = "auto",
    ):
        """
        Convert content (HTML or Markdown) to a formatted DOCX.
        
        Args:
            content: The input content (HTML or Markdown)
            output_filepath: Optional path to save the document
            reference_docx_path: Optional reference document for styles
            input_format: "html", "markdown", or "auto" (default: auto-detect)
        
        Returns:
            True if saved to file, or BytesIO buffer if no output path specified
        """
        if input_format == "auto":
            # Simple auto-detection: if content starts with HTML tags, treat as HTML
            stripped = content.strip()
            if stripped.startswith("<") and ">" in stripped[:100]:
                input_format = "html"
            else:
                input_format = "markdown"
        
        if input_format == "html":
            return self.html_to_docx(content, output_filepath, reference_docx_path)
        else:
            return self.markdown_to_docx(content, output_filepath, reference_docx_path)

    # Convenience wrappers -------------------------------------------------------

    def markdown_to_memory(self, markdown_text: str, reference_docx_path: str | None = None):
        """Convert markdown text to a Word document in memory and return binary data."""
        result = self.markdown_to_docx(markdown_text, None, reference_docx_path)
        
        # Check if result is a BytesIO object (success case)
        if isinstance(result, io.BytesIO):
            # Return the binary content
            return result.getvalue()
        
        # If we got a boolean or other type, something went wrong
        raise ValueError(f"Document conversion failed: received error: {result}")

    def html_to_memory(self, html_text: str, reference_docx_path: str | None = None):
        """Convert HTML text to a Word document in memory and return binary data."""
        result = self.html_to_docx(html_text, None, reference_docx_path)
        
        # Check if result is a BytesIO object (success case)
        if isinstance(result, io.BytesIO):
            # Return the binary content
            return result.getvalue()
        
        # If we got a boolean or other type, something went wrong
        raise ValueError(f"Document conversion failed: received error: {result}")

    def to_memory(self, content: str, reference_docx_path: str | None = None, input_format: str = "auto"):
        """
        Convert content (HTML or Markdown) to a Word document in memory and return binary data.
        
        Args:
            content: The input content (HTML or Markdown)
            reference_docx_path: Optional reference document for styles
            input_format: "html", "markdown", or "auto" (default: auto-detect)
        
        Returns:
            Binary data of the Word document
        """
        result = self.to_docx(content, None, reference_docx_path, input_format)
        
        # Check if result is a BytesIO object (success case)
        if isinstance(result, io.BytesIO):
            # Return the binary content
            return result.getvalue()
        
        # If we got a boolean or other type, something went wrong
        raise ValueError(f"Document conversion failed: received error: {result}")

    # Runtime setters ------------------------------------------------------------

    def set_font(self, font_name: str, font_size: int | None = None):
        self.font_name = font_name
        if font_size is not None:
            self.font_size = font_size

    def set_heading_style(
        self,
        font: str | None = None,
        sizes: dict | None = None,
        color: str | None = None,
    ):
        if font is not None:
            self.heading_font = font
        if sizes is not None:
            self.heading_sizes = sizes
        if color is not None:
            self.heading_color = color

    def set_page_orientation(self, orientation: str):
        if orientation.lower() in {"portrait", "landscape"}:
            self.page_orientation = orientation.lower()

    def set_margins(self, margins: dict):
        if isinstance(margins, dict) and all(k in margins for k in ("top", "bottom", "left", "right")):
            self.margins = margins